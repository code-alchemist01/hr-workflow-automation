# -*- coding: utf-8 -*-

# --- GEREKLİ KÜTÜPHANELER ---
import os
import json
import requests
import spacy
import cv2
# pytesseract kaldırıldı
from moviepy.editor import VideoFileClip
from docx import Document
from elevenlabs import ElevenLabs
from langchain.text_splitter import RecursiveCharacterTextSplitter
from deepface import DeepFace
from collections import Counter
import warnings
import numpy as np
from PIL import Image
import base64
import google.generativeai as genai
import io
from datetime import datetime

# DeepFace ve TensorFlow uyarılarını gizle
warnings.filterwarnings("ignore", category=UserWarning)


# Tesseract OCR kaldırıldı 

# --- KONFİGÜRASYON ---
ELEVENLABS_API_KEY = "your_elevenlabs_api_key_here"
GEMINI_API_KEY = "your_gemini_api_key_here"  
INPUT_VIDEO_FILE = "video1723838072.mp4"
TEMP_AUDIO_FILE = "gecici_ses.wav"
TRANSCRIPT_DOCX_FILE = "mulakat_transkripti.docx"
SANIYEDE_ANALIZ_SAYISI = 2
TEMP_FRAME_FILE = "temp_frame.jpg"
TEMP_FACE_FILE = "temp_face.jpg"
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"
MODEL_NAME = "qwen/qwen3-4b-2507"
FINAL_ANALYSIS_TXT_FILE = "analiz_sonucu.docx"
DURATION_THRESHOLD = 20
    
# DeepFace duygu etiketlerini Türkçeye çevirme sözlüğü
DUYGU_SOZLUGU = {
    'angry': 'ÖFKELİ',
    'disgust': 'TİKSİNMİŞ',
    'fear': 'KORKMUŞ',
    'happy': 'MUTLU',
    'sad': 'ÜZGÜN',
    'surprise': 'ŞAŞKIN',
    'neutral': 'DOĞAL'
}

# --- YENİ PROMPT'LAR ---
PROMPT_SCORING_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

BÖLÜM 1: PUANLAMA TABLOSU
• Her kriteri, madde imi (•) ile başlayan ayrı bir satırda ve şu formatta yaz:
• Kriter Adı: (Puan/5) - {candidate_name}'in [puanın nedenini açıklayan kısa ve tanımlayıcı bir cümle].
• Değerlendirilecek Kriterler:
• İletişim Becerisi
• Motivasyon ve Tutku
• Kültürel Uyum
• Analitik/Düşünsel Beceriler
• Profesyonel Tutum
• Geçmiş Deneyim Uyumu
• Liderlik ve Girişimcilik
• Zayıflıklarla Başa Çıkma Yetisi
• Uzun Vadeli Potansiyel
• Genel Etki / İzlenim
• Analiz Sonu:
• Genel Ortalama Puan: Tüm puanların ortalamasını, ondalık ayraç olarak virgül kullanarak hesapla. Örnek: (3,86/5)
• İK Genel Yorum: Adayın genel potansiyelini ve ana bulguları özetleyen birkaç cümlelik bir paragraf yaz."""

PROMPT_RECRUITER_DETAILS = """
Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.

BÖLÜM 2: RECRUITER NOTU TALİMATLARI

• Analizini, aşağıdaki altı başlığın tamamını madde imi (•) ile başlayan ayrı satırlar olarak yapılandır.
• Her başlığın altına, mülakat özetinden çıkardığın somut bilgilere (projeler, deneyimler, yetenekler) dayanarak detaylı ve profesyonel bir metin yaz. Adayın adını metin içinde uygun yerlerde kullan.
• Başlıklar ve İçerikleri:
• Aday Adı: Mülakat metninden adayın adını çıkar.
• Pozisyon: Mülakat metninden adayın başvurduğu pozisyonu belirle.
• Genel Yorum: Adayın geçmişi, deneyim süresi ve genel performansı hakkında özet bir paragraf yaz.
• Dikkat Çeken Güçlü Yönler: Adayın öne çıkan teknik veya sosyal yeteneklerini, projelerden örnekler vererek anlat.
• Geliştirme Alanları: Adayın hangi konularda kendini geliştirebileceğini ve potansiyel gelişim alanlarını belirt.
• Değerlendirme Önerisi: Aday için bir sonraki adımları (ikinci görüşme, teknik test vb.) ve gelişimini destekleyecek önerileri (eğitim, kurs vb.) içeren bir paragraf yaz.
"""

PROMPT_QA_MATCHING = """
Aşağıdaki mülakat özetini analiz ederek soru-cevap eşleştirmesi yap.

BÖLÜM 3: SORU-CEVAP ANALİZİ

• Mülakat metninden sorulan soruları ve verilen cevapları tespit et.
• Her soru-cevap çifti için aşağıdaki formatı kullan:
• Soru Kategorisi: [Teknik/Davranışsal/Genel]
• Soru: [Sorulan soru]
• Cevap Kalitesi: (Puan/5) - [Cevabın detay seviyesi, doğruluğu ve profesyonelliği]
• Cevap Özeti: [Adayın verdiği cevabın kısa özeti]
• Eksik Noktalar: [Cevaplamadığı veya yetersiz kaldığı konular]

• Analiz Sonu:
• Toplam Soru Sayısı: [Sayı]
• En İyi Cevaplanan Sorular: [2-3 soru kategorisi]
• Gelişim Gerektiren Alanlar: [Zayıf cevaplanan konular]
"""

PROMPT_TECHNICAL_COMPETENCY = """
Aşağıdaki mülakat özetini teknik yetkinlik açısından analiz et.

BÖLÜM 4: TEKNİK YETKİNLİK DEĞERLENDİRMESİ

• Adayın bahsettiği teknik konuları kategorize et:
• Programlama Dilleri: [Bahsedilen diller ve deneyim seviyeleri]
• Teknolojiler/Framework'ler: [Kullandığı teknolojiler]
• Projeler: [Bahsettiği projeler ve teknik detayları]
• Problem Çözme: [Teknik problemlere yaklaşımı]
• Öğrenme Yeteneği: [Yeni teknolojileri öğrenme konusundaki tutumu]

• Her kategori için puanlama:
• Kategori Adı: (Puan/5) - [Yetkinlik seviyesi ve gerekçesi]

• Teknik Analiz Özeti:
• Güçlü Olduğu Teknik Alanlar: [En iyi olduğu 2-3 alan]
• Gelişim Alanları: [Eksik veya zayıf olduğu teknik konular]
• Önerilen Teknik Eğitimler: [Hangi konularda eğitim alması önerilir]
"""

PROMPT_SOFT_SKILLS = """
Aşağıdaki mülakat özetini soft skill (yumuşak beceri) açısından analiz et.

BÖLÜM 5: SOFT SKİLL ANALİZİ

• Mülakat boyunca gözlemlenen soft skill'leri değerlendir:
• İletişim Tarzı: [Açık, net, etkili iletişim kurma yeteneği]
• Takım Çalışması: [Ekip içinde çalışma deneyimi ve yaklaşımı]
• Liderlik: [Liderlik deneyimi ve potansiyeli]
• Adaptasyon: [Değişime uyum sağlama yeteneği]
• Zaman Yönetimi: [Proje ve görev yönetimi becerileri]
• Stres Yönetimi: [Baskı altında çalışma yeteneği]
• Yaratıcılık: [Yenilikçi çözümler üretme yeteneği]
• Empati: [Başkalarını anlama ve işbirliği kurma]

• Her soft skill için puanlama:
• Beceri Adı: (Puan/5) - [Gözlemlenen davranış örnekleri]

• Soft Skill Özeti:
• En Güçlü Soft Skill'ler: [En gelişmiş 3 beceri]
• Gelişim Alanları: [Güçlendirilmesi gereken beceriler]
• Kişilik Profili: [Genel kişilik özelliklerinin özeti]
• Takım Uyumu: [Hangi tür takımlarda daha başarılı olabileceği]
"""

# ==============================================================================
# --- YARDIMCI FONKSİYONLAR ---
# ==============================================================================

def extract_candidate_name_from_text(transcript_file_path):
    """
    Transkript dosyasından adayın ismini LLM ile çıkarır.
    """
    try:
        interview_text = read_text_from_docx(transcript_file_path)
        if not interview_text or not interview_text.strip():
            print(f"HATA: '{transcript_file_path}' dosyası bulunamadı veya boş.")
            return None

        prompt_name = f"""
        Aşağıdaki mülakat metninden adayın ismini çıkar. Sadece ismi döndür, başka bir şey yazma.
        Örneğin: "Merhaba, ben Can Bey" -> Can Bey

        --- MÜLAKAT METNİ ---
        {interview_text}
        --- İSİM ---
        """
        candidate_name = get_llm_analysis(prompt_name, MODEL_NAME)
        if candidate_name and candidate_name.strip():
            print(f"Metinden tespit edilen isim: {candidate_name}")
            return candidate_name.strip()
        else:
            print("Uyarı: Metinden isim tespit edilemedi.")
            return None
    except Exception as e:
        print(f"HATA: Metinden isim çıkarılırken hata: {e}")
        return None

def extract_frame_and_name(video_path, text_name, max_duration=120.0):
    """
    Gelişmiş yüz tespit algoritması ile videodan en iyi yüz görüntüsünü çıkarır.
    Çoklu yüz tespit yöntemi ve kalite kontrolü içerir.
    """
    try:
        video = cv2.VideoCapture(video_path)
        if not video.isOpened():
            print(f"HATA: '{video_path}' videosu açılamadı.")
            return None, None
        
        fps = video.get(cv2.CAP_PROP_FPS) or 30
        max_frames = int(max_duration * fps)
        frame_interval = int(fps * 1)  # Her saniyede bir kare (daha sık kontrol)
        frame_num = 0
        final_name = text_name
        best_face_data = None
        best_face_score = 0

        print(f"Gelişmiş yüz tespit başlatılıyor... (İlk {max_duration} saniye taranacak)")

        while frame_num < max_frames:
            ret, frame = video.read()
            if not ret:
                break

            if frame_num % frame_interval == 0:
                timestamp = frame_num / fps
                
                # Çoklu yüz tespit yöntemi
                face_data = detect_best_face_in_frame(frame, timestamp)
                
                if face_data and face_data['quality_score'] > best_face_score:
                    best_face_data = face_data
                    best_face_score = face_data['quality_score']
                    print(f"Daha iyi yüz bulundu (skor: {best_face_score:.2f}, zaman: {timestamp:.1f}s)")

            frame_num += 1

        video.release()
        
        # En iyi yüzü kaydet
        face_image_path = None
        if best_face_data:
            face_image_path = save_best_face(best_face_data)
            print(f"En iyi yüz görüntüsü kaydedildi: {face_image_path}")
            print(f"Kalite skoru: {best_face_score:.2f}")
        else:
            print("Uyarı: Hiçbir yüz tespit edilemedi.")

        return final_name, face_image_path
    except Exception as e:
        print(f"HATA: Gelişmiş yüz tespit sırasında hata: {e}")
        return None, None

def detect_best_face_in_frame(frame, timestamp):
    """
    Bir karede en iyi yüzü tespit eder ve kalite skorunu hesaplar.
    """
    try:
        # DeepFace ile yüz tespiti
        analysis = DeepFace.analyze(frame, actions=['emotion'], enforce_detection=False, silent=True)
        
        if isinstance(analysis, list) and len(analysis) > 0:
            face_info = analysis[0]
            face_region = face_info['region']
            x, y, w, h = face_region['x'], face_region['y'], face_region['w'], face_region['h']
            
            # Yüz görüntüsünü çıkar
            face_img = frame[y:y+h, x:x+w]
            
            # Kalite kontrolü
            quality_score = calculate_face_quality(face_img, w, h)
            
            return {
                'face_image': face_img,
                'region': face_region,
                'quality_score': quality_score,
                'timestamp': timestamp,
                'emotion_data': face_info.get('emotion', {})
            }
    except Exception:
        pass
    
    return None

def calculate_face_quality(face_img, width, height):
    """
    Yüz görüntüsünün kalitesini hesaplar (0-100 arası skor).
    """
    try:
        # Boyut skoru (büyük yüzler daha iyi)
        size_score = min(100, (width * height) / 10000 * 100)
        
        # Netlik skoru (Laplacian varyansı)
        gray = cv2.cvtColor(face_img, cv2.COLOR_BGR2GRAY)
        laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
        sharpness_score = min(100, laplacian_var / 500 * 100)
        
        # Parlaklık skoru (çok karanlık veya çok aydınlık değil)
        brightness = np.mean(gray)
        brightness_score = 100 - abs(brightness - 128) / 128 * 100
        
        # Toplam kalite skoru
        total_score = (size_score * 0.4 + sharpness_score * 0.4 + brightness_score * 0.2)
        
        return total_score
    except Exception:
        return 0

def save_best_face(face_data):
    """
    En iyi yüz görüntüsünü kaydeder ve dosya yolunu döndürür.
    """
    try:
        face_img = face_data['face_image']
        
        # Görüntü iyileştirme
        enhanced_face = enhance_face_image(face_img)
        
        # Dosyayı kaydet
        cv2.imwrite(TEMP_FACE_FILE, enhanced_face)
        
        return TEMP_FACE_FILE
    except Exception as e:
        print(f"Yüz kaydetme hatası: {e}")
        return None

def enhance_face_image(face_img):
    """
    Yüz görüntüsünü iyileştirir (kontrast, parlaklık, netlik).
    """
    try:
        # Histogram eşitleme
        lab = cv2.cvtColor(face_img, cv2.COLOR_BGR2LAB)
        lab[:,:,0] = cv2.equalizeHist(lab[:,:,0])
        enhanced = cv2.cvtColor(lab, cv2.COLOR_LAB2BGR)
        
        # Hafif keskinleştirme
        kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        sharpened = cv2.filter2D(enhanced, -1, kernel)
        
        # Orijinal ile karışım (çok agresif olmaması için)
        result = cv2.addWeighted(enhanced, 0.7, sharpened, 0.3, 0)
        
        return result
    except Exception:
        return face_img

# compare_names fonksiyonu kaldırıldı (OCR ile birlikte)

def analyze_character_from_image(face_image_path):
    """
    Gelişmiş Gemini 2.5 Flash ile adayın yüz görüntüsüne dayalı kapsamlı karakter analizi yapar.
    """
    try:
        print(f"Gelişmiş görsel analiz başlatılıyor: '{face_image_path}'")

        # Gemini API'sini yapılandır
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash")

        # Görseli base64 formatına çevir
        image_base64 = image_to_base64(face_image_path)
        if image_base64 is None:
            print("Görsel yüklenemedi. Analiz atlanıyor.")
            return None

        # Gelişmiş analiz prompt'u
        prompt = """
        Bu görsel bir mülakat adayının profesyonel ortamdaki görüntüsüdür. Lütfen aşağıdaki kategorilerde detaylı analiz yapın:

        🎯 **PROFESYONEL GÖRÜNÜM ANALİZİ**
        
        **1. Kıyafet ve Grooming (25 puan)**
        - Kıyafet seçimi ve uygunluğu (iş ortamına uygun mu?)
        - Temizlik ve düzenlilik
        - Renk uyumu ve stil
        - Saç düzeni ve genel bakım
        - Aksesuar kullanımı (varsa)
        
        **2. Beden Dili ve Postür (25 puan)**
        - Oturuş/duruş pozisyonu
        - Omuz hizası ve sırt düzlüğü
        - El pozisyonları ve jestler
        - Genel vücut dili (açık/kapalı)
        - Kendine güven yansıması
        
        **3. Yüz İfadesi ve Göz Kontaktu (25 puan)**
        - Yüz ifadesinin genel tonu
        - Göz kontaktu kalitesi
        - Gülümseme ve mimikler
        - Stres/gerginlik belirtileri
        - Odaklanma ve dikkat
        
        **4. Genel İlk İzlenim (25 puan)**
        - Profesyonellik düzeyi
        - Hazırlık ve özen gösterme
        - Özgüven ve kararlılık
        - İş ortamına uyum potansiyeli
        - Güvenilirlik hissi
        
        📊 **PUANLAMA SİSTEMİ**
        Her kategori için 0-25 puan verin ve toplam 100 üzerinden değerlendirin.
        
        📝 **RAPOR FORMATI**
        Her kategori için:
        - Gözlemlenen özellikler
        - Puan ve gerekçesi
        - Öneriler (varsa)
        
        ⚠️ **ÖNEMLI NOTLAR**
        - Sadece görsel olarak gözlemlenebilir unsurları değerlendirin
        - Önyargısız ve objektif olun
        - Her değerlendirmenin görsel kanıtını belirtin
        - Kültürel farklılıkları göz önünde bulundurun
        
        Lütfen analizi Türkçe olarak yapın ve profesyonel bir dil kullanın.
        """

        # Görseli ve prompt'u Gemini API'sine gönder
        response = model.generate_content([
            prompt,
            {
                "mime_type": "image/jpeg",
                "data": image_base64
            }
        ])
        
        analysis = response.text
        print("Gelişmiş görsel analiz başarıyla tamamlandı.")
        
        # Analiz sonucunu formatla
        formatted_analysis = format_visual_analysis(analysis)
        
        return formatted_analysis
    except Exception as e:
        print(f"HATA: Gelişmiş görüntü analizi sırasında hata: {e}")
        return None

def image_to_base64(image_path):
    """
    Görüntü dosyasını base64 formatına çevirir.
    """
    try:
        with Image.open(image_path) as img:
            # Görüntü boyutunu optimize et (çok büyükse küçült)
            max_size = (800, 800)
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            buffered = io.BytesIO()
            img.save(buffered, format="JPEG", quality=85)
            return base64.b64encode(buffered.getvalue()).decode("utf-8")
    except FileNotFoundError:
        print(f"Hata: Görsel dosyası bulunamadı: {image_path}")
        return None
    except Exception as e:
        print(f"Görsel işleme hatası: {e}")
        return None

def format_visual_analysis(analysis):
    """
    Görsel analiz sonucunu formatlar ve yapılandırır.
    """
    try:
        # Analiz başlığı ekle
        formatted = "\n" + "="*60 + "\n"
        formatted += "           GÖRSEL TABALLI PROFESYONEL ANALİZ\n"
        formatted += "="*60 + "\n\n"
        
        # Ana analizi ekle
        formatted += analysis
        
        # Analiz tarihi ekle
        from datetime import datetime
        formatted += "\n\n" + "-"*40
        formatted += f"\nAnaliz Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        formatted += "Analiz Yöntemi: Gemini 2.5 Flash - Gelişmiş Görsel AI\n"
        
        return formatted
    except Exception:
        return analysis

def videodaki_duygulari_analiz_et(video_path, saniyede_kontrol=2):
    """Gelişmiş video duygu analizi - optimize edilmiş frame işleme ve duygu tespiti."""
    print("\n🎭 Gelişmiş video duygu analizi başlatılıyor...")
    
    video = cv2.VideoCapture(video_path)
    if not video.isOpened():
        print("❌ Video dosyası açılamadı!")
        return []
    
    # Video bilgilerini al
    fps = video.get(cv2.CAP_PROP_FPS)
    total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = total_frames / fps if fps > 0 else 0
    
    print(f"📹 Video Bilgileri: {duration:.1f}s, {fps:.1f} FPS, {total_frames} frame")
    
    frame_interval = max(1, int(fps / saniyede_kontrol)) if saniyede_kontrol > 0 and fps > 0 else int(fps)
    
    duygu_zaman_cizelgesi = []
    frame_num = 0
    processed_frames = 0
    failed_detections = 0
    
    # İlerleme takibi
    progress_interval = max(1, total_frames // 20)  # %5'lik aralıklarla ilerleme
    
    while video.isOpened():
        ret, frame = video.read()
        if not ret:
            break

        if frame_num % frame_interval == 0:
            try:
                # Frame kalitesini kontrol et
                if is_frame_suitable_for_emotion_analysis(frame):
                    analysis = DeepFace.analyze(
                        frame,
                        actions=['emotion'],
                        enforce_detection=False,
                        silent=True
                    )
                    
                    if isinstance(analysis, list) and len(analysis) > 0:
                        emotion_data = analysis[0]
                        dominant_emotion_en = emotion_data['dominant_emotion']
                        dominant_emotion_tr = DUYGU_SOZLUGU.get(dominant_emotion_en, dominant_emotion_en.upper())
                        
                        # Duygu güven skorunu al
                        confidence = emotion_data['emotion'].get(dominant_emotion_en, 0)
                        
                        timestamp = frame_num / fps
                        
                        duygu_entry = {
                            'zaman': timestamp,
                            'duygu': dominant_emotion_tr,
                            'guven_skoru': confidence,
                            'tum_duygular': {DUYGU_SOZLUGU.get(k, k): v for k, v in emotion_data['emotion'].items()}
                        }
                        
                        duygu_zaman_cizelgesi.append(duygu_entry)
                        processed_frames += 1
                    else:
                        failed_detections += 1
                else:
                    failed_detections += 1
                    
            except Exception as e:
                failed_detections += 1
                if frame_num % (frame_interval * 10) == 0:  # Her 10 frame'de bir hata logla
                    print(f"⚠️ Frame {frame_num} analiz hatası: {str(e)[:50]}...")
        
        # İlerleme göster
        if frame_num % progress_interval == 0:
            progress = (frame_num / total_frames) * 100
            print(f"📊 İlerleme: {progress:.1f}% ({processed_frames} başarılı, {failed_detections} başarısız)")
        
        frame_num += 1

    video.release()
    
    # Sonuç özeti
    print(f"\n✅ Duygu analizi tamamlandı!")
    print(f"📈 Toplam analiz edilen frame: {processed_frames}")
    print(f"❌ Başarısız tespit: {failed_detections}")
    print(f"🎯 Başarı oranı: {(processed_frames/(processed_frames+failed_detections)*100):.1f}%" if (processed_frames+failed_detections) > 0 else "🎯 Başarı oranı: 0%")
    
    # Duygu dağılımını analiz et
    if duygu_zaman_cizelgesi:
        emotion_summary = analyze_emotion_distribution(duygu_zaman_cizelgesi)
        print(f"🎭 Dominant duygu: {emotion_summary['dominant']} ({emotion_summary['dominant_percentage']:.1f}%)")
    
    return duygu_zaman_cizelgesi

def is_frame_suitable_for_emotion_analysis(frame):
    """Frame'in duygu analizi için uygun olup olmadığını kontrol eder."""
    try:
        # Frame boyutunu kontrol et
        if frame is None or frame.size == 0:
            return False
        
        height, width = frame.shape[:2]
        if height < 100 or width < 100:
            return False
        
        # Parlaklık kontrolü (çok karanlık veya çok parlak frameler)
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        mean_brightness = np.mean(gray)
        
        # Çok karanlık (< 30) veya çok parlak (> 220) frameler
        if mean_brightness < 30 or mean_brightness > 220:
            return False
        
        # Bulanıklık kontrolü (Laplacian variance)
        laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
        if laplacian_var < 100:  # Çok bulanık
            return False
        
        return True
    except Exception:
        return False

def analyze_emotion_distribution(emotion_timeline):
    """Duygu zaman çizelgesini analiz eder ve dağılım bilgisi verir."""
    try:
        emotions = [entry['duygu'] for entry in emotion_timeline]
        emotion_counts = Counter(emotions)
        total_count = len(emotions)
        
        if total_count == 0:
            return {'dominant': 'BELİRSİZ', 'dominant_percentage': 0, 'distribution': {}}
        
        dominant_emotion = emotion_counts.most_common(1)[0][0]
        dominant_count = emotion_counts[dominant_emotion]
        dominant_percentage = (dominant_count / total_count) * 100
        
        distribution = {emotion: (count / total_count) * 100 for emotion, count in emotion_counts.items()}
        
        return {
            'dominant': dominant_emotion,
            'dominant_percentage': dominant_percentage,
            'distribution': distribution,
            'total_analyzed_frames': total_count
        }
    except Exception:
        return {'dominant': 'BELİRSİZ', 'dominant_percentage': 0, 'distribution': {}}

def assess_video_quality(video_path):
    """Video kalitesini kapsamlı olarak değerlendirir."""
    print("\n🔍 Video kalitesi değerlendiriliyor...")
    
    try:
        video = cv2.VideoCapture(video_path)
        if not video.isOpened():
            return {'overall_score': 0, 'issues': ['Video açılamadı'], 'recommendations': ['Video dosyasını kontrol edin']}
        
        # Video özelliklerini al
        fps = video.get(cv2.CAP_PROP_FPS)
        width = int(video.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(video.get(cv2.CAP_PROP_FRAME_HEIGHT))
        total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = total_frames / fps if fps > 0 else 0
        
        quality_metrics = {
            'resolution_score': 0,
            'fps_score': 0,
            'duration_score': 0,
            'brightness_score': 0,
            'sharpness_score': 0,
            'stability_score': 0
        }
        
        issues = []
        recommendations = []
        
        # Çözünürlük değerlendirmesi
        total_pixels = width * height
        if total_pixels >= 1920 * 1080:  # Full HD+
            quality_metrics['resolution_score'] = 100
        elif total_pixels >= 1280 * 720:  # HD
            quality_metrics['resolution_score'] = 80
        elif total_pixels >= 854 * 480:  # SD
            quality_metrics['resolution_score'] = 60
        else:
            quality_metrics['resolution_score'] = 30
            issues.append(f"Düşük çözünürlük: {width}x{height}")
            recommendations.append("En az 720p (1280x720) çözünürlük kullanın")
        
        # FPS değerlendirmesi
        if fps >= 30:
            quality_metrics['fps_score'] = 100
        elif fps >= 24:
            quality_metrics['fps_score'] = 80
        elif fps >= 15:
            quality_metrics['fps_score'] = 60
        else:
            quality_metrics['fps_score'] = 30
            issues.append(f"Düşük FPS: {fps:.1f}")
            recommendations.append("En az 24 FPS kullanın")
        
        # Süre değerlendirmesi
        if 60 <= duration <= 1800:  # 1-30 dakika ideal
            quality_metrics['duration_score'] = 100
        elif 30 <= duration <= 3600:  # 30 saniye - 1 saat kabul edilebilir
            quality_metrics['duration_score'] = 80
        else:
            quality_metrics['duration_score'] = 50
            if duration < 30:
                issues.append(f"Çok kısa video: {duration:.1f}s")
                recommendations.append("En az 30 saniye video kaydedin")
            else:
                issues.append(f"Çok uzun video: {duration/60:.1f} dakika")
                recommendations.append("Video süresini 30 dakika altında tutun")
        
        # Frame kalitesi analizi (örnekleme ile)
        sample_frames = min(50, total_frames // 10)  # En fazla 50 frame örnekle
        frame_interval = max(1, total_frames // sample_frames)
        
        brightness_scores = []
        sharpness_scores = []
        
        for i in range(0, total_frames, frame_interval):
            video.set(cv2.CAP_PROP_POS_FRAMES, i)
            ret, frame = video.read()
            if not ret:
                break
            
            # Parlaklık analizi
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            brightness = np.mean(gray)
            
            # Optimal parlaklık aralığı: 80-180
            if 80 <= brightness <= 180:
                brightness_score = 100
            elif 50 <= brightness <= 220:
                brightness_score = 70
            else:
                brightness_score = 30
            brightness_scores.append(brightness_score)
            
            # Keskinlik analizi (Laplacian variance)
            laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
            if laplacian_var > 500:
                sharpness_score = 100
            elif laplacian_var > 200:
                sharpness_score = 80
            elif laplacian_var > 100:
                sharpness_score = 60
            else:
                sharpness_score = 30
            sharpness_scores.append(sharpness_score)
        
        video.release()
        
        # Ortalama skorları hesapla
        quality_metrics['brightness_score'] = np.mean(brightness_scores) if brightness_scores else 0
        quality_metrics['sharpness_score'] = np.mean(sharpness_scores) if sharpness_scores else 0
        quality_metrics['stability_score'] = 90  # Basit varsayım, geliştirilebilir
        
        # Kalite sorunlarını tespit et
        if quality_metrics['brightness_score'] < 70:
            issues.append("Parlaklık sorunları tespit edildi")
            recommendations.append("Daha iyi aydınlatma kullanın")
        
        if quality_metrics['sharpness_score'] < 70:
            issues.append("Bulanıklık/odak sorunları tespit edildi")
            recommendations.append("Kamerayı sabit tutun ve odağı kontrol edin")
        
        # Genel skor hesapla
        overall_score = np.mean(list(quality_metrics.values()))
        
        quality_assessment = {
            'overall_score': round(overall_score, 1),
            'metrics': quality_metrics,
            'video_info': {
                'resolution': f"{width}x{height}",
                'fps': fps,
                'duration': duration,
                'total_frames': total_frames
            },
            'issues': issues,
            'recommendations': recommendations,
            'quality_level': get_quality_level(overall_score)
        }
        
        print(f"📊 Video kalitesi: {quality_assessment['quality_level']} ({overall_score:.1f}/100)")
        
        return quality_assessment
        
    except Exception as e:
        print(f"❌ Video kalitesi değerlendirme hatası: {e}")
        return {
            'overall_score': 0,
            'issues': [f'Değerlendirme hatası: {str(e)}']
        }

def get_quality_level(score):
    """Skor bazında kalite seviyesi döndürür."""
    if score >= 90:
        return "Mükemmel"
    elif score >= 80:
        return "Çok İyi"
    elif score >= 70:
        return "İyi"
    elif score >= 60:
        return "Orta"
    elif score >= 50:
        return "Zayıf"
    else:
        return "Çok Zayıf"

def assess_image_quality(image_path):
    """Görüntü kalitesini değerlendirir."""
    try:
        img = cv2.imread(image_path)
        if img is None:
            return {'score': 0, 'issues': ['Görüntü yüklenemedi']}
        
        height, width = img.shape[:2]
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Çözünürlük skoru
        total_pixels = width * height
        if total_pixels >= 500000:  # ~700x700+
            resolution_score = 100
        elif total_pixels >= 200000:  # ~450x450+
            resolution_score = 80
        else:
            resolution_score = 50
        
        # Parlaklık skoru
        brightness = np.mean(gray)
        if 80 <= brightness <= 180:
            brightness_score = 100
        elif 50 <= brightness <= 220:
            brightness_score = 70
        else:
            brightness_score = 30
        
        # Keskinlik skoru
        laplacian_var = cv2.Laplacian(gray, cv2.CV_64F).var()
        if laplacian_var > 500:
            sharpness_score = 100
        elif laplacian_var > 200:
            sharpness_score = 80
        else:
            sharpness_score = 50
        
        overall_score = (resolution_score + brightness_score + sharpness_score) / 3
        
        issues = []
        if resolution_score < 80:
            issues.append("Düşük çözünürlük")
        if brightness_score < 70:
            issues.append("Parlaklık problemi")
        if sharpness_score < 70:
            issues.append("Bulanıklık")
        
        return {
            'score': round(overall_score, 1),
            'resolution_score': resolution_score,
            'brightness_score': brightness_score,
            'sharpness_score': sharpness_score,
            'dimensions': f"{width}x{height}",
            'issues': issues,
            'quality_level': get_quality_level(overall_score)
        }
        
    except Exception as e:
        return {'score': 0, 'issues': [f'Değerlendirme hatası: {str(e)}']}

def konusmacilari_ayir_ve_cumlele(words_data, duygu_cizelgesi):
    """
    ElevenLabs'ten gelen kelime bazlı diarizasyon verisini işler.
    Konuşmacı etiketlerine göre metni gruplar, spaCy ile cümlelere ayırır ve duygu analizi ekler.
    """
    if not words_data:
        print("\nKonuşmacı ayırmak için kelime verisi bulunamadı.")
        return []

    speaker_map = {}
    next_speaker_num = 1
    current_speaker_id_raw = None
    current_speaker_mapped_label = None
    current_speaker_words_buffer = []
    all_speaker_utterances_combined = []

    for word_obj in words_data:
        word_text = word_obj.text
        raw_speaker_id = word_obj.speaker_id
        start_time = word_obj.start
        end_time = word_obj.end

        if raw_speaker_id not in speaker_map:
            speaker_map[raw_speaker_id] = f"Konuşmacı {next_speaker_num}"
            next_speaker_num += 1

        mapped_speaker_label = speaker_map[raw_speaker_id]

        if current_speaker_id_raw is None:
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
        elif raw_speaker_id != current_speaker_id_raw:
            if current_speaker_words_buffer:
                all_speaker_utterances_combined.append({
                    'speaker': current_speaker_mapped_label,
                    'text_combined': " ".join([w['text'] for w in current_speaker_words_buffer]).strip(),
                    'start_time': current_speaker_words_buffer[0]['start'],
                    'end_time': current_speaker_words_buffer[-1]['end']
                })
            current_speaker_id_raw = raw_speaker_id
            current_speaker_mapped_label = mapped_speaker_label
            current_speaker_words_buffer = []

        current_speaker_words_buffer.append({'text': word_text, 'start': start_time, 'end': end_time})

    if current_speaker_words_buffer:
        all_speaker_utterances_combined.append({
            'speaker': current_speaker_mapped_label,
            'text_combined': " ".join([w['text'] for w in current_speaker_words_buffer]).strip(),
            'start_time': current_speaker_words_buffer[0]['start'],
            'end_time': current_speaker_words_buffer[-1]['end']
        })

    final_diarized_sentences = []
    try:
        nlp = spacy.load("xx_ent_wiki_sm")
        if "sentencizer" not in nlp.pipe_names:
            nlp.add_pipe("sentencizer")
    except OSError:
        print("\nspaCy modeli 'xx_ent_wiki_sm' bulunamadı.")
        print("Lütfen 'python -m spacy download xx_ent_wiki_sm' komutunu çalıştırın.")
        return []
    except Exception as e:
        print(f"spaCy modeli yüklenirken bir hata oluştu: {e}")
        return []

    for entry in all_speaker_utterances_combined:
        doc = nlp(entry['text_combined'])
        sentences_for_speaker = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        start_time = entry['start_time']
        end_time = entry['end_time']
        duration = end_time - start_time
        sentences_count = len(sentences_for_speaker)

        if sentences_count == 0:
            continue

        for i, sent in enumerate(sentences_for_speaker):
            sent_start = start_time + (i * duration / sentences_count)
            sent_end = start_time + ((i + 1) * duration / sentences_count)
            ilgili_duygular = [d['duygu'] for d in duygu_cizelgesi if sent_start <= d['zaman'] <= sent_end]
            dominant_duygu = Counter(ilgili_duygular).most_common(1)[0][0] if ilgili_duygular else "BELİRSİZ"

            final_diarized_sentences.append({
                "konusmaci": entry['speaker'],
                "diyalog": sent,
                "duygu": dominant_duygu,
                "baslangic": sent_start,
                "bitis": sent_end
            })

    return final_diarized_sentences

def read_text_from_docx(file_path):
    """
    Bir .docx dosyasındaki tüm metni okur ve tek bir metin bloğu olarak döndürür.
    """
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        print(f"HATA: '{file_path}' dosyası okunurken hata oluştu: {e}")
        return None

def get_llm_analysis(prompt, model_name):
    """
    LM Studio API aracılığıyla yerel LLM'e bir prompt gönderir ve analizi alır.
    """
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 2048
    }
    try:
        response = requests.post(LM_STUDIO_API_URL, json=payload)
        response.raise_for_status()
        response_data = response.json()
        if "choices" in response_data and response_data["choices"]:
            return response_data["choices"][0]["message"]["content"]
        else:
            print(f"HATA: API yanıtında 'choices' anahtarı bulunamadı veya boş. Yanıt: {response_data}")
            return None
    except requests.exceptions.ConnectionError:
        print("HATA: LM Studio API sunucusuna bağlanılamadı. Lütfen sunucunun çalıştığından emin olun.")
        return None
    except requests.exceptions.HTTPError as http_err:
        print(f"HATA: API isteği sırasında HTTP hatası: {http_err}\nYanıt İçeriği: {response.text}")
        return None
    except Exception as e:
        print(f"HATA: Analiz sırasında beklenmedik bir hata oluştu: {e}")
        return None

def write_analysis_to_txt(file_path, analysis_scoring, analysis_recruiter, image_analysis, candidate_name, analysis_qa=None, analysis_technical=None, analysis_soft_skills=None):
    """
    Analiz sonuçlarını yeni bir .docx dosyasına yazar, düzgün formatlanmış şekilde.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Yeni bir Word belgesi oluştur
        doc = Document()
        
        # Başlık
        title = doc.add_heading('MÜLAKAT ANALİZİ SONUCU', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('=' * 50)
        doc.add_paragraph()
        
        # 1. Aday Değerlendirme Puanlama Tablosu
        heading1 = doc.add_heading('1. ADAY DEĞERLENDİRME PUANLAMA TABLOSU', level=2)
        doc.add_paragraph('-' * 45)
        doc.add_paragraph(analysis_scoring)
        doc.add_paragraph()
        
        # 2. Recruiter Notu
        heading2 = doc.add_heading('2. RECRUITER NOTU', level=2)
        doc.add_paragraph('-' * 17)
        doc.add_paragraph(analysis_recruiter)
        doc.add_paragraph()
        
        # 3. Görüntü Tabanlı Karakter Analizi
        heading3 = doc.add_heading('3. GÖRÜNTÜ TABANLI KARAKTER ANALİZİ', level=2)
        doc.add_paragraph('-' * 35)
        if image_analysis:
            doc.add_paragraph(image_analysis)
        else:
            doc.add_paragraph(f"{candidate_name} için görüntü tabanlı analiz yapılamadı.")
        doc.add_paragraph()
        
        # 4. Soru-Cevap Analizi
        if analysis_qa:
            heading4 = doc.add_heading('4. SORU-CEVAP ANALİZİ', level=2)
            doc.add_paragraph('-' * 22)
            doc.add_paragraph(analysis_qa)
            doc.add_paragraph()
        
        # 5. Teknik Yetkinlik Değerlendirmesi
        if analysis_technical:
            heading5 = doc.add_heading('5. TEKNİK YETKİNLİK DEĞERLENDİRMESİ', level=2)
            doc.add_paragraph('-' * 37)
            doc.add_paragraph(analysis_technical)
            doc.add_paragraph()
        
        # 6. Soft Skill Analizi
        if analysis_soft_skills:
            heading6 = doc.add_heading('6. SOFT SKILL ANALİZİ', level=2)
            doc.add_paragraph('-' * 22)
            doc.add_paragraph(analysis_soft_skills)
        
        # Dosyayı kaydet
        doc.save(file_path)
        print(f"\nAnaliz sonuçları '{file_path}' dosyasına başarıyla kaydedildi.")
        
    except Exception as e:
        print(f"HATA: Sonuçlar dosyaya yazılırken hata oluştu: {e}")
        # Hata durumunda eski txt formatına geri dön
        try:
            with open(file_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                f.write('MÜLAKAT ANALİZİ SONUCU\n')
                f.write('=' * 50 + '\n\n')
                
                f.write('1. ADAY DEĞERLENDİRME PUANLAMA TABLOSU\n')
                f.write('-' * 45 + '\n')
                f.write(analysis_scoring + '\n\n')
                
                f.write('2. RECRUITER NOTU\n')
                f.write('-' * 17 + '\n')
                f.write(analysis_recruiter + '\n\n')
                
                f.write('3. GÖRÜNTÜ TABANLI KARAKTER ANALİZİ\n')
                f.write('-' * 35 + '\n')
                if image_analysis:
                    f.write(image_analysis + '\n\n')
                else:
                    f.write(f"{candidate_name} için görüntü tabanlı analiz yapılamadı.\n\n")
                
                if analysis_qa:
                    f.write('4. SORU-CEVAP ANALİZİ\n')
                    f.write('-' * 22 + '\n')
                    f.write(analysis_qa + '\n\n')
                
                if analysis_technical:
                    f.write('5. TEKNİK YETKİNLİK DEĞERLENDİRMESİ\n')
                    f.write('-' * 37 + '\n')
                    f.write(analysis_technical + '\n\n')
                
                if analysis_soft_skills:
                    f.write('6. SOFT SKILL ANALİZİ\n')
                    f.write('-' * 22 + '\n')
                    f.write(analysis_soft_skills + '\n')
                    
            print(f"\nYedek olarak TXT formatında kaydedildi.")
        except Exception as txt_error:
            print(f"HATA: TXT formatında da kayıt yapılamadı: {txt_error}")

def adim_2_metin_analizi_yap(transcript_file_path, candidate_name="Aday"):
    """
    ADIM 2: Transkripti DOCX dosyasından okur, LLM'e analiz için gönderir ve sonuçları döndürür.
    """
    print(f"--- ADIM 2: '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return None, None

    print("Mülakat transkripti başarıyla okundu.")

    print(f"'{MODEL_NAME}' modeline 'Aday Değerlendirme Puanlama Tablosu' için istek gönderiliyor...")
    prompt_scoring = PROMPT_SCORING_DETAILS.format(candidate_name=candidate_name)
    prompt_scoring = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et. 
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {prompt_scoring}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama analizi alınamadı. İşlem durduruluyor.")
        return None, None
    print("Puanlama analizi başarıyla tamamlandı.")

    print(f"'{MODEL_NAME}' modeline 'Recruiter Notu' için istek gönderiliyor...")
    prompt_recruiter = PROMPT_RECRUITER_DETAILS.format(candidate_name=candidate_name)
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {prompt_recruiter}

    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi alınamadı. İşlem durduruluyor.")
        return None, None
    print("Recruiter notu analizi başarıyla tamamlandı.")

    return analysis_scoring, analysis_recruiter

def adim_2_metin_analizi_chunk(transcript_file_path, candidate_name="Aday"):
    """
    ADIM 2 (Chunk Tabanlı): Transkripti DOCX dosyasından okur, metni parçalara böler,
    her parça için özet oluşturur, özetleri birleştirir ve tek bir nihai analiz yapar.
    """
    print(f"--- ADIM 2 (Chunk Tabanlı): '{transcript_file_path}' Dosyasındaki Metin Analiz Ediliyor ---")

    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Analiz yapılamıyor.")
        return None, None

    print("Mülakat transkripti başarıyla okundu.")

    print("Metin parçalara ayrılıyor...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(interview_text)
    print(f"Metin {len(chunks)} parçaya bölündü.")

    chunk_summaries = []
    print("Her bir metin parçası için özetler oluşturuluyor...")
    for i, chunk in enumerate(chunks):
        print(f"Parça {i + 1}/{len(chunks)} işleniyor...")
        prompt_chunk_summary = f"""
        Aşağıdaki mülakat metni parçasını oku. Bu parçadan, adayın aşağıda listelenen yetkinlikleri ile ilgili
        tüm önemli bilgileri, kilit ifadeleri ve somut örnekleri Türkçe olarak maddeler halinde özetle.
        Sadece bu metin parçasında geçen bilgileri kullan.

        Yetkinlikler:
        - İletişim Becerisi
        - Motivasyon ve Tutku
        - Kültürel Uyum
        - Analitik/Düşünsel Beceriler
        - Profesyonel Tutum
        - Geçmiş Deneyim Uyumu
        - Liderlik ve Girişimcilik
        - Zayıflıklarla Başa Çıkma Yetisi
        - Uzun Vadeli Potansiyel
        - Genel Etki / İzlenim

        --- MÜLAKAT METNİ PARÇASI ---
        {chunk}
        --- ÖZETİNİ BURAYA BAŞLAT ---
        """
        summary = get_llm_analysis(prompt_chunk_summary, MODEL_NAME)
        if summary:
            chunk_summaries.append(summary)
        else:
            print(f"Parça {i + 1} için özet alınamadı. İşlem durduruluyor.")
            return None, None

    combined_summary = "\n\n---\n\n".join(chunk_summaries)
    print("\nTüm parçaların özetleri başarıyla birleştirildi. Nihai analiz başlıyor.")

    print("Nihai 'Aday Değerlendirme Puanlama Tablosu' oluşturuluyor...")
    prompt_scoring = PROMPT_SCORING_DETAILS.format(candidate_name=candidate_name)
    prompt_scoring = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "1. Aday Değerlendirme Puanlama Tablosu" başlığı altında yapılandır ve detaylandır:
    {prompt_scoring}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_scoring = get_llm_analysis(prompt_scoring, MODEL_NAME)
    if not analysis_scoring:
        print("Puanlama tablosu analizi oluşturulamadı. İşlem durduruluyor.")
        return None, None

    print("Nihai 'Recruiter Notu' oluşturuluyor...")
    prompt_recruiter = PROMPT_RECRUITER_DETAILS.format(candidate_name=candidate_name)
    prompt_recruiter = f"""
    Aşağıdaki mülakat özetini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece "2. Recruiter Notu" başlığı altında yapılandır ve detaylandır:
    {prompt_recruiter}

    --- MÜLAKAT ÖZETİ METNİ ---
    {combined_summary}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_recruiter = get_llm_analysis(prompt_recruiter, MODEL_NAME)
    if not analysis_recruiter:
        print("Recruiter notu analizi oluşturulamadı. İşlem durduruluyor.")
        return None, None

    return analysis_scoring, analysis_recruiter

def advanced_content_analysis(transcript_file_path, candidate_name="Aday"):
    """
    Gelişmiş içerik analizi: Q&A Matching, Technical Competency, Soft Skills
    """
    print(f"--- GELİŞMİŞ İÇERİK ANALİZİ: '{transcript_file_path}' Dosyası Analiz Ediliyor ---")
    
    interview_text = read_text_from_docx(transcript_file_path)
    if not interview_text or not interview_text.strip():
        print(f"HATA: '{transcript_file_path}' dosyası bulunamadı, bozuk veya boş. Gelişmiş analiz yapılamıyor.")
        return None, None, None
    
    print("Mülakat transkripti başarıyla okundu. Gelişmiş analizler başlıyor...")
    
    # Q&A Matching Analizi
    print(f"'{MODEL_NAME}' modeline 'Soru-Cevap Analizi' için istek gönderiliyor...")
    prompt_qa = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {PROMPT_QA_MATCHING}
    
    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_qa = get_llm_analysis(prompt_qa, MODEL_NAME)
    if not analysis_qa:
        print("Soru-Cevap analizi alınamadı.")
        analysis_qa = "Soru-Cevap analizi yapılamadı."
    else:
        print("Soru-Cevap analizi başarıyla tamamlandı.")
    
    # Technical Competency Analizi
    print(f"'{MODEL_NAME}' modeline 'Teknik Yetkinlik Değerlendirmesi' için istek gönderiliyor...")
    prompt_tech = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {PROMPT_TECHNICAL_COMPETENCY}
    
    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_technical = get_llm_analysis(prompt_tech, MODEL_NAME)
    if not analysis_technical:
        print("Teknik Yetkinlik analizi alınamadı.")
        analysis_technical = "Teknik Yetkinlik analizi yapılamadı."
    else:
        print("Teknik Yetkinlik analizi başarıyla tamamlandı.")
    
    # Soft Skills Analizi
    print(f"'{MODEL_NAME}' modeline 'Soft Skill Analizi' için istek gönderiliyor...")
    prompt_soft = f"""
    Aşağıdaki mülakat metnini profesyonel bir İnsan Kaynakları (İK) uzmanı gibi analiz et.
    Analizinin tamamını, mülakat metninin dili ne olursa olsun, mutlaka Türkçe olarak oluştur.
    Analizini sadece aşağıdaki başlık altında yapılandır ve detaylandır:
    {PROMPT_SOFT_SKILLS}
    
    --- MÜLAKAT METNİ ---
    {interview_text}
    --- ANALİZİNİ BURAYA BAŞLAT ---
    """
    analysis_soft_skills = get_llm_analysis(prompt_soft, MODEL_NAME)
    if not analysis_soft_skills:
        print("Soft Skill analizi alınamadı.")
        analysis_soft_skills = "Soft Skill analizi yapılamadı."
    else:
        print("Soft Skill analizi başarıyla tamamlandı.")
    
    return analysis_qa, analysis_technical, analysis_soft_skills

def extract_questions(interview_text):
    """
    Mülakat metninden soruları çıkarır
    """
    questions = []
    lines = interview_text.split('\n')
    
    for line in lines:
        line = line.strip()
        # Soru işaretli cümleleri tespit et
        if '?' in line and len(line) > 10:
            # Konuşmacı bilgisini temizle
            if ']:' in line:
                question = line.split(']:')[-1].strip()
            else:
                question = line
            
            if question and question not in questions:
                questions.append(question)
    
    return questions

def categorize_question(question):
    """
    Soruyu kategorize eder: Teknik, Davranışsal, Genel
    """
    question_lower = question.lower()
    
    # Teknik sorular için anahtar kelimeler
    technical_keywords = ['kod', 'programlama', 'teknoloji', 'yazılım', 'algoritma', 
                         'veri', 'database', 'api', 'framework', 'python', 'java',
                         'makine öğrenmesi', 'yapay zeka', 'ai', 'ml']
    
    # Davranışsal sorular için anahtar kelimeler
    behavioral_keywords = ['deneyim', 'proje', 'takım', 'liderlik', 'çatışma', 
                          'zorluk', 'başarı', 'hata', 'öğrenme', 'gelişim',
                          'motivasyon', 'hedef', 'çalışma tarzı']
    
    # Teknik kategori kontrolü
    for keyword in technical_keywords:
        if keyword in question_lower:
            return "Teknik"
    
    # Davranışsal kategori kontrolü
    for keyword in behavioral_keywords:
        if keyword in question_lower:
            return "Davranışsal"
    
    # Varsayılan olarak Genel
    return "Genel"

# ==============================================================================
# --- ANA AKIŞ FONKSİYONLARI ---
# ==============================================================================

def adim_1_videodan_metne_cevir():
    """
    ADIM 1: Videodan sesi çıkarır, ElevenLabs ile deşifre eder, duygu analizi yapar ve sonucu DOCX'e kaydeder.
    """
    print("--- ADIM 1: Mülakat Videosu Metne Çevriliyor ve Duygu Analizi Yapılıyor ---")

    # 1.1. Video dosyasının varlığını kontrol et
    if not os.path.exists(INPUT_VIDEO_FILE):
        print(f"HATA: '{INPUT_VIDEO_FILE}' video dosyası bulunamadı. Lütfen kontrol edin.")
        return None, None, None

    # 1.2. Videodan sesi çıkar ve süreyi hesapla
    try:
        print(f"'{INPUT_VIDEO_FILE}' videosundan ses çıkarılıyor...")
        with VideoFileClip(INPUT_VIDEO_FILE) as video:
            video_duration_minutes = video.duration / 60
            video.audio.write_audiofile(TEMP_AUDIO_FILE)
        print(f"Ses başarıyla '{TEMP_AUDIO_FILE}' olarak kaydedildi.")
        print(f"Video süresi: {video_duration_minutes:.2f} dakika")
    except Exception as e:
        print(f"HATA: Video işlenirken bir hata oluştu: {e}")
        return None, None, None

    # 1.3. ElevenLabs ile sesi yazıya çevir
    words_data = []
    try:
        print(f"'{TEMP_AUDIO_FILE}' dosyası ElevenLabs API'sine gönderiliyor (diarize=True)...")
        elevenlabs_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)
        with open(TEMP_AUDIO_FILE, "rb") as audio_file:
            response = elevenlabs_client.speech_to_text.convert(
                file=audio_file,
                model_id="scribe_v1",
                diarize=True,
            )
            words_data = response.words
        print("Ses, ElevenLabs tarafından başarıyla deşifre edildi.")
    except Exception as e:
        print(f"HATA: ElevenLabs STT işlemi sırasında hata: {e}")
        return None, None, None
    finally:
        if os.path.exists(TEMP_AUDIO_FILE):
            os.remove(TEMP_AUDIO_FILE)
            print(f"Geçici ses dosyası '{TEMP_AUDIO_FILE}' silindi.")

    # 1.4. Videodan duygu analizi yap
    duygu_cizelgesi = videodaki_duygulari_analiz_et(INPUT_VIDEO_FILE, SANIYEDE_ANALIZ_SAYISI)
    if not duygu_cizelgesi:
        print("Uyarı: Duygu analizi yapılamadı. Transkript duygu bilgisi olmadan devam edecek.")

    # 1.5. Metni konuşmacılara göre ayır ve duygu bilgisi ekle
    if not words_data:
        print("HATA: ElevenLabs'tan kelime bazlı veri alınamadı. İşlem durduruluyor.")
        return None, None, None

    diarized_output = konusmacilari_ayir_ve_cumlele(words_data, duygu_cizelgesi)
    if not diarized_output:
        print("HATA: Konuşmacı ayrımı yapılamadı veya boş sonuç döndü.")
        return None, None, None

    # 1.6. Sonucu Word dosyasına yaz
    try:
        doc = Document()
        doc.add_heading('Mülakat Transkripti', level=1)
        for entry in diarized_output:
            start_min = int(entry['baslangic'] // 60)
            start_sec = int(entry['baslangic'] % 60)
            line = f"[{entry['konusmaci']}][{start_min}:{start_sec:02d}][{entry['duygu']}]: {entry['diyalog']}"
            doc.add_paragraph(line)
        doc.save(TRANSCRIPT_DOCX_FILE)
        print(f"Transkript başarıyla '{TRANSCRIPT_DOCX_FILE}' dosyasına kaydedildi.")
        print("--- ADIM 1 TAMAMLANDI ---\n")
        return TRANSCRIPT_DOCX_FILE, video_duration_minutes, duygu_cizelgesi
    except Exception as e:
        print(f"HATA: Transkript Word dosyasına yazılırken hata oluştu: {e}")
        return None, None, None

# ==============================================================================
# --- BETİĞİ ÇALIŞTIR ---
# ==============================================================================

if __name__ == "__main__":
    print("🎯 ===== GELİŞMİŞ MÜLAKAT VİDEOSU ANALİZ SÜRECİ BAŞLATILDI =====")
    print("🔧 Yeni Özellikler: Gelişmiş Yüz Tespiti | Akıllı Görsel Analiz | Video Kalite Kontrolü")
    print("="*80)

    # Adım 0: Video kalitesi ön değerlendirmesi
    print("\n📋 ADIM 0: Video Kalitesi Ön Değerlendirmesi")
    video_quality = assess_video_quality(INPUT_VIDEO_FILE)
    
    if video_quality['overall_score'] < 50:
        print(f"⚠️ UYARI: Video kalitesi düşük ({video_quality['overall_score']}/100)")
        print("❌ Tespit edilen sorunlar:")
        for issue in video_quality.get('issues', []):
            print(f"   • {issue}")
        print("💡 Öneriler:")
        for rec in video_quality.get('recommendations', []):
            print(f"   • {rec}")
        
        user_choice = input("\n❓ Düşük kaliteli video ile devam etmek istiyor musunuz? (e/h): ").lower()
        if user_choice != 'e':
            print("🛑 İşlem kullanıcı tarafından durduruldu.")
            exit(1)
    else:
        print(f"✅ Video kalitesi: {video_quality['quality_level']} ({video_quality['overall_score']}/100)")

    # Adım 1: Videodan transkript ve gelişmiş duygu analizi
    print("\n📋 ADIM 1: Video İşleme ve Gelişmiş Duygu Analizi")
    transcript_file, video_duration, duygu_cizelgesi = adim_1_videodan_metne_cevir()

    if transcript_file and video_duration is not None:
        # Adım 2: Transkriptten adayın ismini çıkar
        print("\n📋 ADIM 2: Aday İsmi Tespiti")
        candidate_name = extract_candidate_name_from_text(transcript_file)
        if not candidate_name:
            candidate_name = "Aday"
            print("⚠️ Uyarı: Metinden isim tespit edilemedi, varsayılan isim 'Aday' kullanılacak.")
        else:
            print(f"✅ Aday ismi tespit edildi: {candidate_name}")

        # Adım 3: Gelişmiş yüz tespiti ve görüntü çıkarma
        print("\n📋 ADIM 3: Gelişmiş Yüz Tespiti ve Görüntü Çıkarma")
        final_name, face_image_path = extract_frame_and_name(INPUT_VIDEO_FILE, candidate_name, max_duration=60.0)

        # Adım 4: Kapsamlı görüntü tabanlı karakter analizi
        print("\n📋 ADIM 4: Kapsamlı Görüntü Tabanlı Karakter Analizi")
        image_analysis = None
        image_quality_report = None
        
        if face_image_path and final_name == candidate_name:
            # Çıkarılan görüntünün kalitesini değerlendir
            image_quality_report = assess_image_quality(face_image_path)
            print(f"📊 Çıkarılan görüntü kalitesi: {image_quality_report['quality_level']} ({image_quality_report['score']}/100)")
            
            if image_quality_report['score'] >= 60:
                image_analysis = analyze_character_from_image(face_image_path)
                print("✅ Gelişmiş görüntü tabanlı karakter analizi tamamlandı.")
            else:
                print(f"⚠️ Görüntü kalitesi düşük ({image_quality_report['score']}/100), analiz atlanıyor.")
                print("❌ Tespit edilen sorunlar:", ", ".join(image_quality_report.get('issues', [])))
        else:
            print(f"❌ {candidate_name} için uygun yüz görüntüsü bulunamadı, görüntü analizi atlanıyor.")

        # Adım 5: Metin tabanlı analiz
        print("\n📋 ADIM 5: Metin Tabanlı Analiz")
        if video_duration > DURATION_THRESHOLD:
            print(f"📊 Video süresi {video_duration:.2f} dakika, chunk tabanlı analiz kullanılıyor.")
            analysis_scoring, analysis_recruiter = adim_2_metin_analizi_chunk(transcript_file, candidate_name)
        else:
            print(f"📊 Video süresi {video_duration:.2f} dakika, standart analiz kullanılıyor.")
            analysis_scoring, analysis_recruiter = adim_2_metin_analizi_yap(transcript_file, candidate_name)

        # Adım 6: Gelişmiş içerik analizi
        print("\n📋 ADIM 6: Gelişmiş İçerik Analizi")
        analysis_qa, analysis_technical, analysis_soft_skills = advanced_content_analysis(transcript_file, candidate_name)

        # Adım 7: Kapsamlı rapor oluşturma
        print("\n📋 ADIM 7: Kapsamlı Analiz Raporu Oluşturma")
        if analysis_scoring and analysis_recruiter:
            # Görsel analiz raporunu zenginleştir
            if image_analysis and image_quality_report:
                enhanced_image_analysis = f"""
{image_analysis}

📊 GÖRÜNTÜ KALİTE RAPORU
{"="*40}
Genel Kalite Skoru: {image_quality_report['score']}/100 ({image_quality_report['quality_level']})
Çözünürlük: {image_quality_report['dimensions']}
Çözünürlük Skoru: {image_quality_report['resolution_score']}/100
Parlaklık Skoru: {image_quality_report['brightness_score']}/100
Keskinlik Skoru: {image_quality_report['sharpness_score']}/100

📹 VİDEO KALİTE RAPORU
{"="*40}
Genel Video Kalitesi: {video_quality['overall_score']}/100 ({video_quality['quality_level']})
Çözünürlük: {video_quality['video_info']['resolution']}
FPS: {video_quality['video_info']['fps']:.1f}
Süre: {video_quality['video_info']['duration']:.1f} saniye
"""
            else:
                enhanced_image_analysis = image_analysis
            
            write_analysis_to_txt(
                FINAL_ANALYSIS_TXT_FILE, 
                analysis_scoring, 
                analysis_recruiter, 
                enhanced_image_analysis, 
                candidate_name, 
                analysis_qa, 
                analysis_technical, 
                analysis_soft_skills
            )
            
            print(f"✅ Kapsamlı analiz raporu '{FINAL_ANALYSIS_TXT_FILE}' dosyasına kaydedildi.")
        else:
            print("❌ Metin analizi tamamlanamadı, rapor oluşturulamadı.")
    else:
        print("\n❌ Süreç, video işleme aşamasındaki bir hata nedeniyle durduruldu.")

    # Geçici dosyaları temizle
    print("\n🧹 Geçici Dosyalar Temizleniyor...")
    for temp_file in [TEMP_FRAME_FILE, TEMP_FACE_FILE]:
        if os.path.exists(temp_file):
            os.remove(temp_file)
            print(f"🗑️ Geçici dosya '{temp_file}' silindi.")

    print("\n🎉 ===== GELİŞMİŞ ANALİZ SÜRECİ BAŞARIYLA TAMAMLANDI =====")
    print("📊 Analiz Özeti:")
    print(f"   • Video Kalitesi: {video_quality.get('quality_level', 'Bilinmiyor')}")
    print(f"   • Görsel Analiz: {'✅ Tamamlandı' if image_analysis else '❌ Atlandı'}")
    print(f"   • Duygu Analizi: {'✅ Tamamlandı' if duygu_cizelgesi else '❌ Atlandı'}")
    print(f"   • Metin Analizi: {'✅ Tamamlandı' if analysis_scoring else '❌ Başarısız'}")
    print("🎯 Gelişmiş IIAS sistemi ile analiz tamamlandı!")